import io
import unittest
from unittest.mock import patch

from docx import Document

from app import app


def _docx_bytes() -> bytes:
    document = Document()
    document.core_properties.title = "Docx Probe"
    document.core_properties.author = "Codex QA"
    document.add_heading("Docx Probe", level=1)
    document.add_paragraph("Probe paragraph.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class AppDocxConversionTests(unittest.TestCase):
    def test_analyze_accepts_docx_and_returns_docx_payload(self) -> None:
        client = app.test_client()
        fake_analysis = {
            "source_type": "docx",
            "profile": "docx_reflow",
            "paragraph_count": 12,
            "heading1_count": 1,
            "heading2_count": 2,
            "heading3_count": 0,
            "list_count": 1,
            "table_count": 1,
            "image_count": 0,
            "hyperlink_count": 1,
            "estimated_sections": 1,
            "publication_analysis": {
                "profile": "docx_reflow",
                "confidence": 0.96,
                "has_toc": True,
                "external_tools": {},
                "profile_reason": "DOCX structure detected.",
            },
        }

        with patch("app.analyze_docx", return_value=fake_analysis):
            response = client.post(
                "/analyze",
                data={"file": (io.BytesIO(_docx_bytes()), "sample.docx")},
                content_type="multipart/form-data",
            )

        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        self.assertTrue(payload["success"])
        self.assertEqual(payload["source_type"], "docx")
        self.assertEqual(payload["analysis"]["heading1_count"], 1)

    def test_analyze_accepts_uppercase_docx_extension(self) -> None:
        client = app.test_client()
        fake_analysis = {
            "source_type": "docx",
            "profile": "docx_reflow",
            "paragraph_count": 4,
            "heading1_count": 1,
            "heading2_count": 0,
            "heading3_count": 0,
            "list_count": 0,
            "table_count": 0,
            "image_count": 0,
            "hyperlink_count": 0,
            "estimated_sections": 1,
            "publication_analysis": {
                "profile": "docx_reflow",
                "confidence": 0.93,
                "has_toc": False,
                "external_tools": {},
                "profile_reason": "DOCX structure detected.",
            },
        }

        with patch("app.analyze_docx", return_value=fake_analysis):
            response = client.post(
                "/analyze",
                data={"file": (io.BytesIO(_docx_bytes()), "SAMPLE.DOCX")},
                content_type="multipart/form-data",
            )

        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        self.assertTrue(payload["success"])
        self.assertEqual(payload["source_type"], "docx")

    def test_convert_accepts_docx_and_sets_source_header(self) -> None:
        client = app.test_client()
        fake_result = {
            "epub_bytes": b"epub-docx",
            "source_type": "docx",
            "analysis": {"profile": "docx_reflow", "confidence": 0.95},
            "quality_report": {
                "validation_status": "passed",
                "validation_tool": "epubcheck",
                "warnings": [],
                "high_risk_pages": [],
                "high_risk_sections": [],
            },
            "document_summary": {
                "title": "Docx Probe",
                "author": "Codex QA",
                "profile": "docx_reflow",
                "layout_mode": "reflowable",
                "section_count": 2,
                "asset_count": 1,
            },
        }

        with patch("app.convert_document_to_epub_with_report", return_value=fake_result):
            response = client.post(
                "/convert",
                data={
                    "file": (io.BytesIO(_docx_bytes()), "sample.docx"),
                    "profile": "auto-premium",
                    "quality_gate_mode": "off",
                    "ocr": "false",
                    "language": "pl",
                    "heading_repair": "false",
                },
                content_type="multipart/form-data",
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data, b"epub-docx")
        self.assertEqual(response.headers.get("X-Source-Type"), "docx")
        self.assertIsNone(response.headers.get("X-PDF-Type"))

    def test_convert_returns_quality_gate_error_when_core_structure_validation_blocks_output(self) -> None:
        client = app.test_client()
        fake_result = {
            "epub_bytes": b"quality-gate-epub",
            "source_type": "docx",
            "analysis": {"profile": "docx_reflow", "confidence": 0.95},
            "quality_report": {
                "validation_status": "passed",
                "validation_tool": "epubcheck",
                "warnings": [],
                "high_risk_pages": [],
                "high_risk_sections": [],
            },
            "document_summary": {
                "title": "Blocked Sample",
                "author": "KindleMaster QA",
                "profile": "docx_reflow",
                "layout_mode": "reflowable",
                "section_count": 2,
                "asset_count": 1,
            },
        }
        validation_report = {
            "package": {
                "status": "failed",
                "errors": ["Manifest has 1 duplicate id(s)."],
                "warnings": [],
            },
            "internal_links": {
                "status": "passed",
                "errors": [],
                "warnings": [],
            },
            "external_links": {
                "status": "passed",
                "errors": [],
                "warnings": [],
            },
            "metadata": {"title": "Blocked Sample", "creator": "KindleMaster QA", "language": "en"},
            "document_stats": {
                "documents_parsed": 1,
                "documents_with_duplicate_ids": 0,
                "links_checked": 0,
                "external_links_checked": 0,
                "manifest_item_count": 2,
                "manifest_targets_missing_count": 0,
                "manifest_duplicate_id_count": 1,
                "navigation_document_count": 1,
                "spine_item_count": 1,
                "spine_linear_item_count": 1,
                "spine_non_linear_item_count": 0,
                "spine_duplicate_targets": 0,
                "spine_unknown_manifest_references": 0,
                "non_linear_spine_targets": 0,
                "unreachable_non_linear_spine_targets": 0,
            },
            "summary": {"status": "failed", "error_count": 1, "warning_count": 0, "epubcheck_status": "passed"},
        }

        with patch("app.convert_document_to_epub_with_report", return_value=fake_result), patch(
            "epub_validation.validate_epub_bytes",
            return_value=validation_report,
        ):
            response = client.post(
                "/convert",
                data={
                    "file": (io.BytesIO(_docx_bytes()), "sample.docx"),
                    "profile": "auto-premium",
                    "ocr": "false",
                    "language": "en",
                    "heading_repair": "false",
                },
                content_type="multipart/form-data",
            )

        self.assertEqual(response.status_code, 422)
        payload = response.get_json()
        self.assertEqual(payload["success"], False)
        self.assertEqual(payload["error_code"], "conversion_quality_gate_failed")
        self.assertEqual(payload["phase"], "quality_gate")
        self.assertIn("Core EPUB structure blocked conversion", payload["error"])

    def test_convert_allows_quality_warnings_in_draft_mode(self) -> None:
        client = app.test_client()
        fake_result = {
            "epub_bytes": b"quality-gate-epub",
            "source_type": "docx",
            "analysis": {"profile": "docx_reflow", "confidence": 0.95},
            "quality_report": {
                "validation_status": "passed",
                "validation_tool": "epubcheck",
                "warnings": [],
                "high_risk_pages": [],
                "high_risk_sections": [],
            },
            "document_summary": {
                "title": "Warned Sample",
                "author": "KindleMaster QA",
                "profile": "docx_reflow",
                "layout_mode": "reflowable",
                "section_count": 2,
                "asset_count": 1,
            },
        }
        validation_report = {
            "package": {
                "status": "passed",
                "errors": [],
                "warnings": ["Niezgodna notka w pakiecie."],
            },
            "internal_links": {
                "status": "passed",
                "errors": [],
                "warnings": [],
            },
            "external_links": {
                "status": "passed",
                "errors": [],
                "warnings": [],
            },
            "metadata": {"title": "Warned Sample", "creator": "KindleMaster QA", "language": "en"},
            "document_stats": {
                "documents_parsed": 1,
                "documents_with_duplicate_ids": 0,
                "links_checked": 8,
                "external_links_checked": 0,
                "manifest_item_count": 2,
                "manifest_targets_missing_count": 0,
                "manifest_duplicate_id_count": 0,
                "navigation_document_count": 1,
                "spine_item_count": 1,
                "spine_linear_item_count": 1,
                "spine_non_linear_item_count": 0,
                "spine_duplicate_targets": 0,
                "spine_unknown_manifest_references": 0,
                "non_linear_spine_targets": 0,
                "unreachable_non_linear_spine_targets": 0,
                "internal_href_with_fragment_count": 0,
                "internal_href_without_fragment_count": 0,
                "internal_href_missing_document_count": 0,
                "internal_href_missing_fragment_count": 0,
            },
            "summary": {"status": "passed", "error_count": 0, "warning_count": 1, "epubcheck_status": "passed"},
        }

        with patch("app.convert_document_to_epub_with_report", return_value=fake_result), patch(
            "epub_validation.validate_epub_bytes",
            return_value=validation_report,
        ):
            response = client.post(
                "/convert",
                data={
                    "file": (io.BytesIO(_docx_bytes()), "sample.docx"),
                    "profile": "auto-premium",
                    "quality_gate_mode": "draft",
                    "ocr": "false",
                    "language": "en",
                    "heading_repair": "false",
                },
                content_type="multipart/form-data",
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data, b"quality-gate-epub")
        self.assertEqual(response.headers.get("X-Source-Type"), "docx")

    def test_convert_blocks_quality_warnings_in_strict_mode(self) -> None:
        client = app.test_client()
        fake_result = {
            "epub_bytes": b"quality-gate-epub",
            "source_type": "docx",
            "analysis": {"profile": "docx_reflow", "confidence": 0.95},
            "quality_report": {
                "validation_status": "passed",
                "validation_tool": "epubcheck",
                "warnings": [],
                "high_risk_pages": [],
                "high_risk_sections": [],
            },
            "document_summary": {
                "title": "Warned Sample",
                "author": "KindleMaster QA",
                "profile": "docx_reflow",
                "layout_mode": "reflowable",
                "section_count": 2,
                "asset_count": 1,
            },
        }
        validation_report = {
            "package": {
                "status": "passed",
                "errors": [],
                "warnings": ["Niezgodna notka w pakiecie."],
            },
            "internal_links": {
                "status": "passed",
                "errors": [],
                "warnings": [],
            },
            "external_links": {
                "status": "passed",
                "errors": [],
                "warnings": [],
            },
            "metadata": {"title": "Warned Sample", "creator": "KindleMaster QA", "language": "en"},
            "document_stats": {
                "documents_parsed": 1,
                "documents_with_duplicate_ids": 0,
                "links_checked": 8,
                "external_links_checked": 0,
                "manifest_item_count": 2,
                "manifest_targets_missing_count": 0,
                "manifest_duplicate_id_count": 0,
                "navigation_document_count": 1,
                "spine_item_count": 1,
                "spine_linear_item_count": 1,
                "spine_non_linear_item_count": 0,
                "spine_duplicate_targets": 0,
                "spine_unknown_manifest_references": 0,
                "non_linear_spine_targets": 0,
                "unreachable_non_linear_spine_targets": 0,
                "internal_href_with_fragment_count": 0,
                "internal_href_without_fragment_count": 0,
                "internal_href_missing_document_count": 0,
                "internal_href_missing_fragment_count": 0,
            },
            "summary": {"status": "passed", "error_count": 0, "warning_count": 1, "epubcheck_status": "passed"},
        }

        with patch("app.convert_document_to_epub_with_report", return_value=fake_result), patch(
            "epub_validation.validate_epub_bytes",
            return_value=validation_report,
        ):
            response = client.post(
                "/convert",
                data={
                    "file": (io.BytesIO(_docx_bytes()), "sample.docx"),
                    "profile": "auto-premium",
                    "quality_gate_mode": "strict",
                    "ocr": "false",
                    "language": "en",
                    "heading_repair": "false",
                },
                content_type="multipart/form-data",
            )

        self.assertEqual(response.status_code, 422)
        payload = response.get_json()
        self.assertEqual(payload["success"], False)
        self.assertEqual(payload["error_code"], "conversion_quality_gate_failed")
        self.assertEqual(payload["phase"], "quality_gate")
        self.assertEqual(payload["quality_gate_mode"], "strict")
        self.assertTrue(payload.get("validation_details"))


if __name__ == "__main__":
    unittest.main()
