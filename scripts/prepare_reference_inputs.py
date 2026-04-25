from __future__ import annotations

from io import BytesIO
import json
import random
import shutil
import tempfile
import textwrap
from pathlib import Path
from typing import Any
import zipfile

import fitz
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFilter, ImageFont


REFERENCE_CASES = [
    {
        "id": "ocr_probe_pdf",
        "document_class": "ocr_probe",
        "input_type": "pdf",
        "language": "pl",
        "quick_smoke": True,
        "release_strict": False,
        "source": "example/ocr_runtime_probe.pdf",
        "target": "reference_inputs/pdf/ocr_probe.pdf",
        "notes": "Tiny OCR probe for the fastest conversion smoke.",
    },
    {
        "id": "ocr_stress_scan_pdf",
        "document_class": "ocr_stress_scan",
        "input_type": "pdf",
        "language": "pl",
        "quick_smoke": False,
        "generator": "ocr_stress_scan",
        "target": "reference_inputs/pdf/ocr_stress_scan.pdf",
        "notes": "Deterministic scanned PDF with OCR-stressed image-only pages and noisy text blocks.",
    },
    {
        "id": "dense_business_guide_pdf",
        "document_class": "dense_business_guide",
        "input_type": "pdf",
        "language": "en",
        "quick_smoke": False,
        "source": "example/BABOK_Guide_v3_Member.pdf",
        "target": "reference_inputs/pdf/dense_business_guide.pdf",
        "notes": "Dense handbook fixture for heading, TOC, and text-cleanup regression.",
    },
    {
        "id": "diagram_training_book_pdf",
        "document_class": "diagram_training_book",
        "input_type": "pdf",
        "language": "en",
        "quick_smoke": False,
        "source": "example/tactits_sample_80pages.pdf",
        "target": "reference_inputs/pdf/diagram_training_book.pdf",
        "notes": "Training/diagram-heavy PDF for layout-sensitive smoke runs.",
    },
    {
        "id": "magazine_layout_pdf",
        "document_class": "magazine_layout",
        "input_type": "pdf",
        "language": "pl",
        "quick_smoke": False,
        "source": "example/02695ab2e05aab728b4b995caa682f947e8be2c3291ff490579797c5a3cc5e26.pdf",
        "target": "reference_inputs/pdf/magazine_layout.pdf",
        "notes": "Magazine/business-report style PDF with layout-heavy structures.",
    },
    {
        "id": "document_like_report_pdf",
        "document_class": "document-like-report",
        "input_type": "pdf",
        "language": "en",
        "quick_smoke": False,
        "target": "reference_inputs/pdf/document_like_report.pdf",
        "notes": "Generated multi-page document-like PDF fixture for manifest-backed corpus reporting.",
        "generator": "document_like_report",
    },
    {
        "id": "scan_probe_epub",
        "document_class": "scan_probe",
        "input_type": "epub",
        "language": "pl",
        "quick_smoke": True,
        "release_strict": False,
        "source": "example/scan_probe_premium.epub",
        "target": "reference_inputs/epub/scan_probe.epub",
        "notes": "Small EPUB for fast validator and repair smoke.",
    },
    {
        "id": "dense_business_guide_epub",
        "document_class": "dense_business_guide",
        "input_type": "epub",
        "language": "en",
        "quick_smoke": False,
        "source": "example/BABOK_current_output.epub",
        "target": "reference_inputs/epub/dense_business_guide.epub",
        "notes": "Representative generated EPUB for release-audit and validator smoke.",
    },
    {
        "id": "simple_report_docx",
        "document_class": "docx_structured_report",
        "input_type": "docx",
        "language": "pl",
        "quick_smoke": True,
        "target": "reference_inputs/docx/simple_report.docx",
        "notes": "Small heading-driven DOCX fixture for quick conversion smoke.",
        "generator": "simple_report",
    },
    {
        "id": "list_table_image_docx",
        "document_class": "docx_rich_content",
        "input_type": "docx",
        "language": "en",
        "quick_smoke": False,
        "target": "reference_inputs/docx/list_table_image.docx",
        "notes": "DOCX fixture with lists, tables, links, and inline imagery.",
        "generator": "list_table_image",
    },
    {
        "id": "no_heading_document_docx",
        "document_class": "docx_no_h1",
        "input_type": "docx",
        "language": "pl",
        "quick_smoke": False,
        "target": "reference_inputs/docx/no_heading_document.docx",
        "notes": "DOCX without Heading 1 styles to test deterministic fallback sectioning.",
        "generator": "no_heading_document",
    },
]


def prepare_reference_inputs(*, root_dir: str | Path = ".") -> dict[str, Any]:
    resolved_root = Path(root_dir).resolve()
    prepared: list[dict[str, Any]] = []
    for case in REFERENCE_CASES:
        target = resolved_root / case["target"]
        target.parent.mkdir(parents=True, exist_ok=True)

        generator = case.get("generator")
        if generator:
            if case["input_type"] == "docx":
                _generate_docx_fixture(generator, target)
            elif case["input_type"] == "pdf":
                _generate_pdf_fixture(generator, target)
            else:  # pragma: no cover - defensive guard
                raise ValueError(f"Unsupported generated fixture type: {case['input_type']}")
            source_path_label = f"<generated:{case['generator']}>"
        else:
            source = resolved_root / case["source"]
            if source.exists():
                shutil.copy2(source, target)
                source_path_label = case["source"]
            else:
                _generate_source_surrogate_fixture(case, target)
                source_path_label = f"<generated-fallback:{case['id']}>"

        prepared_case = {
            **case,
            "source_path": source_path_label,
            "target_path": case["target"],
            "size_bytes": target.stat().st_size,
        }
        prepared.append(prepared_case)

    manifest = {
        "version": 2,
        "root_dir": ".",
        "cases": prepared,
    }
    manifest_path = resolved_root / "reference_inputs" / "manifest.json"
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest


def _generate_docx_fixture(generator: str, target_path: Path) -> None:
    if generator == "simple_report":
        document = _build_simple_report_docx()
    elif generator == "list_table_image":
        document = _build_list_table_image_docx()
    elif generator == "no_heading_document":
        document = _build_no_heading_docx()
    else:  # pragma: no cover - defensive guard
        raise ValueError(f"Unknown DOCX fixture generator: {generator}")
    document.save(str(target_path))


def _generate_pdf_fixture(generator: str, target_path: Path) -> None:
    if generator == "document_like_report":
        document = _build_document_like_report_pdf()
        try:
            document.save(str(target_path))
        finally:
            document.close()
        return
    if generator == "ocr_stress_scan":
        _build_ocr_stress_scan_pdf(target_path)
        return
    raise ValueError(f"Unknown PDF fixture generator: {generator}")


def _generate_source_surrogate_fixture(case: dict[str, Any], target_path: Path) -> None:
    input_type = str(case.get("input_type", "")).lower()
    if input_type == "pdf":
        _generate_source_surrogate_pdf(case, target_path)
        return
    if input_type == "epub":
        _generate_source_surrogate_epub(case, target_path)
        return
    raise FileNotFoundError(case.get("source", target_path))


def _generate_source_surrogate_pdf(case: dict[str, Any], target_path: Path) -> None:
    document = fitz.open()
    title = _case_title(case)
    document.set_metadata(
        {
            "title": title,
            "author": "KindleMaster CI",
            "subject": f"Generated fallback for {case.get('id', 'reference case')}",
            "creator": "KindleMaster prepare_reference_inputs",
            "producer": "PyMuPDF",
        }
    )
    page_count = 3 if case.get("document_class") in {"dense_business_guide", "diagram_training_book", "magazine_layout"} else 1
    try:
        for index in range(1, page_count + 1):
            page = document.new_page(width=595, height=842)
            _draw_pdf_header(page, "KindleMaster generated fallback")
            page.insert_textbox(
                fitz.Rect(64, 96, 531, 160),
                title if index == 1 else f"{title} - section {index}",
                fontsize=20,
                fontname="helv",
                color=(0.10, 0.12, 0.16),
            )
            page.insert_textbox(
                fitz.Rect(64, 170, 531, 420),
                _surrogate_body_text(case, index=index),
                fontsize=11,
                fontname="helv",
                color=(0.15, 0.17, 0.20),
                lineheight=1.25,
            )
            _draw_pdf_footer(page, index, page_count)
        document.save(str(target_path))
    finally:
        document.close()


def _generate_source_surrogate_epub(case: dict[str, Any], target_path: Path) -> None:
    title = _case_title(case)
    language = str(case.get("language", "en") or "en")
    chapter = f"""<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml">
  <head><title>{_xml_escape(title)}</title></head>
  <body>
    <h1 id="intro">{_xml_escape(title)}</h1>
    <p>{_xml_escape(_surrogate_body_text(case, index=1))}</p>
  </body>
</html>
"""
    nav = """<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops">
  <head><title>Navigation</title></head>
  <body><nav epub:type="toc" id="toc"><ol><li><a href="chapter_001.xhtml#intro">Intro</a></li></ol></nav></body>
</html>
"""
    opf = f"""<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="book-id">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:identifier id="book-id">kindlemaster-generated-{_slug(str(case.get('id', 'fixture')))}</dc:identifier>
    <dc:title>{_xml_escape(title)}</dc:title>
    <dc:creator>KindleMaster CI</dc:creator>
    <dc:language>{_xml_escape(language)}</dc:language>
  </metadata>
  <manifest>
    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
    <item id="chapter-001" href="chapter_001.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine>
    <itemref idref="chapter-001"/>
  </spine>
</package>
"""
    target_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(target_path, "w") as archive:
        archive.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        archive.writestr(
            "META-INF/container.xml",
            """<?xml version="1.0" encoding="utf-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="EPUB/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>
""",
        )
        archive.writestr("EPUB/content.opf", opf)
        archive.writestr("EPUB/nav.xhtml", nav)
        archive.writestr("EPUB/chapter_001.xhtml", chapter)


def _case_title(case: dict[str, Any]) -> str:
    label = str(case.get("id") or case.get("document_class") or "reference input").replace("_", " ")
    return label.title()


def _surrogate_body_text(case: dict[str, Any], *, index: int) -> str:
    document_class = str(case.get("document_class", "reference")).replace("_", " ")
    notes = str(case.get("notes", "") or "Generated fixture for KindleMaster verification.")
    return "\n".join(
        [
            f"This generated fallback covers the {document_class} reference class for clean CI checkouts.",
            notes,
            f"Section {index} includes stable prose, predictable headings, and validation-safe markup.",
            "The full local fixture is used automatically when the original source file is present.",
        ]
    )


def _slug(value: str) -> str:
    return "".join(ch.lower() if ch.isalnum() else "-" for ch in value).strip("-") or "fixture"


def _xml_escape(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )


def _build_document_like_report_pdf() -> fitz.Document:
    document = fitz.open()
    document.set_metadata(
        {
            "title": "Document-Like Report Fixture",
            "author": "Anna Nowak",
            "subject": "Deterministic multi-page document-like PDF fixture",
            "creator": "KindleMaster prepare_reference_inputs",
            "producer": "PyMuPDF",
        }
    )
    pages = [
        (
            "Document-Like Report Fixture",
            "Deterministic multi-page PDF generated by scripts/prepare_reference_inputs.py.",
            [
                "By Anna Nowak",
                "Executive summary",
                "This fixture exercises a document-like PDF shape: multiple pages, ordered sections, structured prose, and a stable reading order.",
                "It is intentionally text-heavy so the corpus smoke sees a representative editorial document rather than a scan or OCR probe.",
                "Primary checks include heading recovery, navigation extraction, metadata preservation, and corpus-wide reporting.",
            ],
        ),
        (
            "Scope and method",
            "The fixture is generated deterministically with standard fonts, fixed margins, and explicit page breaks.",
            [
                "1. Page 1 introduces the report title and the evaluation intent.",
                "2. Page 2 records scope, input assumptions, and the expected conversion profile.",
                "3. Page 3 contains the findings block with compact bullet-style observations.",
                "4. Page 4 closes with appendix notes and a reproducible footer.",
            ],
        ),
        (
            "Findings",
            "The content is written to look like a lightweight internal report.",
            [
                "Heading structure should remain stable after conversion.",
                "Paragraph flow should stay continuous across pages.",
                "The generated EPUB should stay small, deterministic, and free from OCR-specific artefacts.",
                "Report metadata should remain concrete enough for manifest-backed corpus-wide summaries.",
            ],
        ),
        (
            "Appendix",
            "Reference table",
            [
                "Document class: document-like-report",
                "Fixture type: generated PDF",
                "Language: en",
                "Pages: 4",
                "Purpose: corpus smoke, manifest-backed execution, and reporting coverage",
            ],
        ),
    ]

    for index, (heading, intro, body_lines) in enumerate(pages, start=1):
        page = document.new_page(width=595, height=842)
        _draw_pdf_header(page, "KindleMaster reference input")
        cursor_y = 112
        cursor_y = _draw_pdf_heading(page, cursor_y, heading)
        cursor_y = _draw_pdf_paragraph(page, cursor_y, intro, indent=0)
        cursor_y = _draw_pdf_gap(cursor_y, gap=10)
        for line in body_lines:
            cursor_y = _draw_pdf_paragraph(page, cursor_y, line, indent=0)
        _draw_pdf_footer(page, index, len(pages))

    document.set_toc(
        [
            [1, "Executive summary", 1],
            [1, "Scope and method", 2],
            [1, "Findings", 3],
            [1, "Appendix", 4],
        ]
    )

    return document


def _build_ocr_stress_scan_pdf(target_path: Path) -> None:
    page_images = [_build_ocr_stress_scan_page(page_number=index + 1, total_pages=3) for index in range(3)]
    document = fitz.open()
    try:
        for page_image in page_images:
            page = document.new_page(width=595, height=842)
            page.insert_image(page.rect, stream=_image_to_png_bytes(page_image), keep_proportion=False)
        document.set_metadata(
            {
                "title": "OCR stress scan",
                "author": "KindleMaster QA",
                "subject": "Deterministic OCR-stressed scanned PDF fixture",
                "keywords": "KindleMaster,OCR,scan",
                "creator": "KindleMaster prepare_reference_inputs",
                "producer": "PyMuPDF",
            }
        )
        document.save(str(target_path), garbage=4, deflate=True, clean=True, no_new_id=True)
    finally:
        document.close()
        for page_image in page_images:
            page_image.close()


def _build_ocr_stress_scan_page(*, page_number: int, total_pages: int) -> Image.Image:
    rng = random.Random(127_000 + page_number)
    base_width = 960
    base_height = 1360
    base = Image.new("L", (base_width, base_height), 246)
    draw = ImageDraw.Draw(base)
    font = ImageFont.load_default()

    draw.rectangle((0, 0, base_width - 1, 82), fill=228)
    draw.text((36, 24), f"OCR stress scan {page_number}/{total_pages}", fill=25, font=font)
    draw.text((36, 52), "Image-only PDF fixture for OCR, scan detection, and cleanup smoke.", fill=55, font=font)
    draw.text((36, 76), "By Anna Nowak", fill=38, font=font)

    left = 42
    top = 138
    lines = [
        "This page intentionally mixes tiny text, numbers, and uneven spacing.",
        "The content stays deterministic while still looking like a scanned handout.",
        "Repeated lines help exercise OCR on faintly skewed, low-contrast blocks.",
        "Alpha 12 | beta 7 | gamma 4 | delta 9 | epsilon 3",
        "Address: 101 Scan Street, Page City, PL 00-127",
    ]
    for index, line in enumerate(lines):
        draw.text((left, top + index * 24), line, fill=35 if index % 2 == 0 else 50, font=font)

    grid_top = 286
    grid_left = 44
    cell_w = 132
    cell_h = 44
    rows = 8
    cols = 5
    for row in range(rows + 1):
        y = grid_top + row * cell_h
        draw.line((grid_left, y, grid_left + cols * cell_w, y), fill=140, width=1)
    for col in range(cols + 1):
        x = grid_left + col * cell_w
        draw.line((x, grid_top, x, grid_top + rows * cell_h), fill=140, width=1)
    table_values = [
        ["ID", "NAME", "STATE", "PAGES", "FLAG"],
        ["A1", "alpha", "ok", "12", "N"],
        ["B2", "beta", "warn", "18", "Y"],
        ["C3", "gamma", "ok", "7", "N"],
        ["D4", "delta", "review", "21", "Y"],
        ["E5", "epsilon", "ok", "5", "N"],
        ["F6", "zeta", "warn", "14", "Y"],
        ["G7", "eta", "ok", "3", "N"],
    ]
    for row_index, row in enumerate(table_values):
        for col_index, value in enumerate(row):
            draw.text((grid_left + 8 + col_index * cell_w, grid_top + 12 + row_index * cell_h), value, fill=30, font=font)

    footer_top = 680
    for line_index in range(20):
        line_y = footer_top + line_index * 20
        draw.text(
            (42, line_y),
            f"{page_number}.{line_index + 1:02d}  OCR stress note {line_index % 4}  scan fixture  {rng.randint(10, 99)}",
            fill=45 if line_index % 3 else 60,
            font=font,
        )

    for _ in range(1800):
        x = rng.randint(0, base_width - 1)
        y = rng.randint(0, base_height - 1)
        base.putpixel((x, y), rng.randint(188, 242))

    rotated = base.rotate(
        -1.2 if page_number % 2 else 0.9,
        resample=Image.Resampling.BICUBIC,
        expand=True,
        fillcolor=246,
    )
    resized = rotated.resize((1654, 2339), resample=Image.Resampling.LANCZOS)
    return resized.filter(ImageFilter.GaussianBlur(radius=0.35))


def _image_to_png_bytes(image: Image.Image) -> bytes:
    buffer = BytesIO()
    image.save(buffer, format="PNG", optimize=True)
    return buffer.getvalue()


def _draw_pdf_header(page, label: str) -> None:
    page.insert_text((54, 44), label, fontname="helv", fontsize=8.5, fill=(0.35, 0.35, 0.35))
    page.draw_line((54, 54), (541, 54), color=(0.82, 0.82, 0.82), width=0.7)


def _draw_pdf_heading(page, y: float, heading: str) -> float:
    page.insert_text((54, y), heading, fontname="helv", fontsize=21, fill=(0.12, 0.16, 0.22))
    return y + 34


def _draw_pdf_paragraph(page, y: float, text: str, *, indent: float = 0.0) -> float:
    wrapped_lines = textwrap.wrap(text, width=74, break_long_words=False, break_on_hyphens=False) or [""]
    for line in wrapped_lines:
        page.insert_text((54 + indent, y), line, fontname="helv", fontsize=11, fill=(0.12, 0.12, 0.12))
        y += 16
    return y + 3


def _draw_pdf_gap(y: float, *, gap: float) -> float:
    return y + gap


def _draw_pdf_footer(page, page_number: int, total_pages: int) -> None:
    page.draw_line((54, 784), (541, 784), color=(0.82, 0.82, 0.82), width=0.7)
    footer = f"Page {page_number} of {total_pages}"
    page.insert_text((54, 798), footer, fontname="helv", fontsize=8.5, fill=(0.35, 0.35, 0.35))


def _build_simple_report_docx() -> Document:
    document = Document()
    document.core_properties.title = "Raport operacyjny"
    document.core_properties.author = "KindleMaster QA"
    title = document.add_heading("Raport operacyjny", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph("To jest mala, uporzadkowana probka DOCX do szybkiego smoke testu.")
    document.add_heading("Zakres", level=2)
    document.add_paragraph("Dokument ma sprawdzic podstawowe mapowanie headingow, akapitow i metadanych.")
    document.add_heading("Wnioski", level=2)
    paragraph = document.add_paragraph("Pelna analiza znajduje sie pod adresem ")
    _append_hyperlink(paragraph, "https://example.com/report", "https://example.com/report")
    paragraph.add_run(".")
    return document


def _build_list_table_image_docx() -> Document:
    document = Document()
    document.core_properties.title = "DOCX Rich Content Probe"
    document.core_properties.author = "KindleMaster QA"
    document.add_heading("Rich content sample", level=1)
    document.add_paragraph("This fixture exercises semantic lists, tables, hyperlinks, and inline imagery.")
    document.add_heading("Checklist", level=2)
    document.add_paragraph("Preserve unordered lists", style="List Bullet")
    document.add_paragraph("Preserve ordered lists", style="List Number")
    document.add_paragraph("Keep hyperlinks clickable", style="List Bullet")
    document.add_heading("Reference link", level=2)
    paragraph = document.add_paragraph("Canonical source: ")
    _append_hyperlink(paragraph, "https://example.com/source", "Example Source")
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Lists"
    table.cell(1, 1).text = "2"
    table.cell(2, 0).text = "Tables"
    table.cell(2, 1).text = "1"
    document.add_heading("Inline image", level=2)
    with _temporary_demo_image() as image_path:
        document.add_picture(str(image_path), width=Inches(2.3))
    return document


def _build_no_heading_docx() -> Document:
    document = Document()
    document.core_properties.title = "Dokument bez H1"
    document.core_properties.author = "KindleMaster QA"
    intro = document.add_paragraph()
    run = intro.add_run("Dokument bez H1")
    run.bold = True
    run.font.size = Pt(18)
    document.add_paragraph("Ta probka nie ma stylu Heading 1 i wymaga deterministycznego fallbacku sekcji.")
    document.add_paragraph("Pierwszy punkt kontrolny", style="List Bullet")
    document.add_paragraph("Drugi punkt kontrolny", style="List Bullet")
    paragraph = document.add_paragraph("Wsparcie: ")
    _append_hyperlink(paragraph, "https://example.com/support", "https://example.com/support")
    return document


class _temporary_demo_image:
    def __enter__(self) -> Path:
        handle = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        handle.close()
        self.path = Path(handle.name)
        image = Image.new("RGB", (420, 220), "#f5efe6")
        draw = ImageDraw.Draw(image)
        draw.rounded_rectangle((18, 18, 402, 202), radius=18, fill="#ffffff", outline="#d97f4a", width=6)
        draw.text((48, 88), "KindleMaster", fill="#1f2937")
        draw.text((48, 120), "DOCX fixture", fill="#6b7280")
        image.save(self.path, format="PNG")
        return self.path

    def __exit__(self, exc_type, exc, tb) -> None:
        if hasattr(self, "path") and self.path.exists():
            self.path.unlink()


def _append_hyperlink(paragraph, url: str, text: str) -> None:
    relation_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def main() -> int:
    manifest = prepare_reference_inputs()
    print(json.dumps(manifest, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
