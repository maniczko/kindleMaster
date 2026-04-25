from __future__ import annotations

import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from unittest.mock import patch
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image, ImageDraw

from converter import ConversionConfig, convert_document_to_epub_with_report, convert_docx_to_epub_with_report
from docx_conversion import analyze_docx, build_docx_publication_document


def _append_hyperlink(paragraph, url: str, text: str) -> None:
    relation_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


class _DemoImage:
    def __enter__(self) -> Path:
        handle = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        handle.close()
        self.path = Path(handle.name)
        image = Image.new("RGB", (320, 180), "#fdf2e8")
        draw = ImageDraw.Draw(image)
        draw.rectangle((18, 18, 302, 162), outline="#d97706", width=5)
        draw.text((42, 74), "DOCX image", fill="#1f2937")
        image.save(self.path, format="PNG")
        return self.path

    def __exit__(self, exc_type, exc, tb) -> None:
        if hasattr(self, "path") and self.path.exists():
            self.path.unlink()


def _build_rich_docx(path: Path) -> None:
    document = Document()
    document.core_properties.title = "Rich DOCX Probe"
    document.core_properties.author = "Codex QA"
    document.add_heading("Opening chapter", level=1)
    document.add_paragraph("Intro paragraph for the first chapter.")
    document.add_heading("Subsection", level=2)
    document.add_paragraph("Bullet item A", style="List Bullet")
    document.add_paragraph("Bullet item B", style="List Bullet")
    document.add_paragraph("Ordered item 1", style="List Number")
    paragraph = document.add_paragraph("Reference: ")
    _append_hyperlink(paragraph, "https://example.com/resource", "https://example.com/resource")
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "42"
    with _DemoImage() as image_path:
        document.add_picture(str(image_path), width=Inches(2.0))
    document.add_heading("Second chapter", level=1)
    document.add_paragraph("Another chapter paragraph.")
    document.save(str(path))


def _build_no_h1_docx(path: Path) -> None:
    document = Document()
    document.core_properties.title = "Fallback DOCX"
    document.core_properties.author = ""
    intro = document.add_paragraph()
    run = intro.add_run("Fallback DOCX")
    run.bold = True
    document.add_paragraph("This document intentionally has no Heading 1 style.")
    document.add_paragraph("List item", style="List Bullet")
    document.save(str(path))


def _build_table_image_docx(path: Path) -> None:
    document = Document()
    document.core_properties.title = "Table Image Probe"
    document.core_properties.author = "Codex QA"
    document.add_heading("Table image chapter", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Evidence"
    table.cell(1, 0).text = "Inline media"
    with _DemoImage() as image_path:
        table.cell(1, 1).paragraphs[0].add_run().add_picture(str(image_path), width=Inches(1.25))
    document.save(str(path))


class DocxConversionTests(unittest.TestCase):
    def test_analyze_docx_reads_metadata_and_structure(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "rich.docx"
            _build_rich_docx(docx_path)

            analysis = analyze_docx(docx_path)

            self.assertEqual(analysis["source_type"], "docx")
            self.assertEqual(analysis["title"], "Rich DOCX Probe")
            self.assertEqual(analysis["author"], "Codex QA")
            self.assertGreaterEqual(analysis["heading1_count"], 2)
            self.assertGreaterEqual(analysis["heading2_count"], 1)
            self.assertGreaterEqual(analysis["list_count"], 3)
            self.assertGreaterEqual(analysis["table_count"], 1)
            self.assertGreaterEqual(analysis["image_count"], 1)
            self.assertGreaterEqual(analysis["hyperlink_count"], 1)
            self.assertEqual(analysis["publication_analysis"]["profile"], "docx_reflow")

    def test_build_docx_publication_document_preserves_rich_structure(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "rich.docx"
            _build_rich_docx(docx_path)

            document = build_docx_publication_document(docx_path, language="en")

            self.assertEqual(document.profile, "docx_reflow")
            self.assertEqual(document.title, "Rich DOCX Probe")
            self.assertEqual(document.author, "Codex QA")
            self.assertEqual(len(document.sections), 2)
            first_section_html = "\n".join(block.raw_html or "" for block in document.sections[0].blocks)
            self.assertIn("<h2>Subsection</h2>", first_section_html)
            self.assertIn("<ul>", first_section_html)
            self.assertIn("<ol>", first_section_html)
            self.assertIn("<table", first_section_html)
            self.assertIn('href="https://example.com/resource"', first_section_html)
            self.assertTrue(any(asset.get("inline") for asset in document.assets))
            self.assertTrue(any(block.block_type == "figure" for block in document.sections[0].blocks))

    def test_build_docx_publication_document_uses_fallback_section_without_h1(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "fallback.docx"
            _build_no_h1_docx(docx_path)

            document = build_docx_publication_document(docx_path, language="pl")

            self.assertEqual(len(document.sections), 1)
            self.assertEqual(document.sections[0].title, "Fallback DOCX")
            self.assertEqual(document.author, "Unknown")

    def test_build_docx_publication_document_preserves_table_cell_images(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "table-image.docx"
            _build_table_image_docx(docx_path)

            document = build_docx_publication_document(docx_path, language="en")

            self.assertEqual(len(document.sections), 1)
            self.assertEqual(len(document.assets), 1)
            table_html = "\n".join(block.raw_html or "" for block in document.sections[0].blocks)
            self.assertIn("<table>", table_html)
            self.assertIn('src="images/docx_', table_html)
            self.assertEqual(len(document.sections[0].assets), 1)

    def test_convert_docx_to_epub_with_report_returns_epub_and_generic_contract(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "rich.docx"
            _build_rich_docx(docx_path)

            result = convert_docx_to_epub_with_report(
                str(docx_path),
                config=ConversionConfig(language="en", profile="auto-premium"),
                original_filename=docx_path.name,
            )

            self.assertEqual(result["source_type"], "docx")
            self.assertIn("quality_report", result)
            self.assertIn("document_summary", result)
            self.assertGreater(len(result["epub_bytes"]), 0)
            with ZipFile(BytesIO(result["epub_bytes"])) as archive:
                self.assertIn("EPUB/content.opf", archive.namelist())
                self.assertTrue(any(name.endswith("nav.xhtml") for name in archive.namelist()))
                image_entries = [name for name in archive.namelist() if name.startswith("EPUB/images/docx_")]
                self.assertEqual(len(image_entries), 1)
                chapters = [
                    archive.read(name).decode("utf-8")
                    for name in archive.namelist()
                    if name.startswith("EPUB/chapter_") and name.endswith(".xhtml")
                ]
                cover = archive.read("EPUB/cover.xhtml").decode("utf-8")

            self.assertTrue(any("EPUB/images" not in chapter and "images/docx_" in chapter for chapter in chapters))
            self.assertNotIn("images/docx_", cover)

    def test_convert_docx_to_epub_with_report_preserves_table_cell_image_asset(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "table-image.docx"
            _build_table_image_docx(docx_path)

            result = convert_docx_to_epub_with_report(
                str(docx_path),
                config=ConversionConfig(language="en", profile="auto-premium"),
                original_filename=docx_path.name,
            )

            with ZipFile(BytesIO(result["epub_bytes"])) as archive:
                image_entries = [name for name in archive.namelist() if name.startswith("EPUB/images/docx_")]
                self.assertEqual(len(image_entries), 1)
                chapter_markup = "\n".join(
                    archive.read(name).decode("utf-8")
                    for name in archive.namelist()
                    if name.startswith("EPUB/chapter_") and name.endswith(".xhtml")
                )

            self.assertIn("<table", chapter_markup)
            self.assertIn("images/docx_", chapter_markup)

    def test_generic_dispatcher_routes_docx(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "rich.docx"
            _build_rich_docx(docx_path)

            result = convert_document_to_epub_with_report(
                str(docx_path),
                config=ConversionConfig(language="en", profile="auto-premium"),
            )

            self.assertEqual(result["source_type"], "docx")

    def test_generic_dispatcher_routes_pdf_to_existing_path(self) -> None:
        with patch("converter.convert_pdf_to_epub_with_report", return_value={"epub_bytes": b"x", "source_type": "pdf"}) as mock_pdf:
            result = convert_document_to_epub_with_report("sample.pdf", config=ConversionConfig())
        self.assertEqual(result["source_type"], "pdf")
        mock_pdf.assert_called_once()


if __name__ == "__main__":
    unittest.main()
